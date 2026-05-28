import os
import re
import base64
from datetime import datetime
from html import escape as html_escape
from io import BytesIO
from typing import Any, Dict


REPORT_SECTION_ORDER = [
    ("report_section_1", "주요 증상"),
    ("report_section_2", "위험 요인"),
    ("report_section_3", "개선 요인"),
    ("report_section_4", "상담사 개입 요인"),
    ("report_section_5", "다음 회기 계획 추천"),
]

DEFAULT_CHART_PLACEHOLDERS = [
    "AI 판별 요약 차트 영역",
    "HIRA 입내원정보 차트 영역",
    "주요 28요인 차트 영역",
]

DEFAULT_CAUTION_TEXT = "AI가 생성한 보고서입니다. 상담사의 전문적 판단에 따라 내용을 검토·수정하여 사용하시기 바랍니다."


def build_report_text(analysis_result: Dict[str, Any]) -> str:
    classification = analysis_result["classification"]
    summary = analysis_result["summary"]

    return f"""[상담보고서 초안]

0. AI 판별 결과
- 우울 관련 라벨: {classification.get("depression", 0)}
- 불안 관련 라벨: {classification.get("anxiety", 0)}
- 중독 관련 라벨: {classification.get("addiction", 0)}

주의:
위 값은 모델 출력 기반 참고값이며, 임상 진단 또는 표준화 검사 점수로 단정하지 않는다.

{summary}
"""


def _parse_markdown_sections(report_text: str) -> Dict[str, str]:
    parsed: Dict[str, str] = {}
    title_to_key = {title: key for key, title in REPORT_SECTION_ORDER}
    pattern = r"^##\s+(.+?)\s*$"
    matches = list(re.finditer(pattern, str(report_text or ""), flags=re.M))

    for index, match in enumerate(matches):
        title = match.group(1).strip()
        key = title_to_key.get(title)

        if not key:
            continue

        start = match.end()
        end = matches[index + 1].start() if index + 1 < len(matches) else len(report_text)
        parsed[key] = str(report_text[start:end]).strip()

    return parsed


def _normalize_report_payload(
    report_text: str,
    metadata: Dict[str, Any] | None = None,
    sections: Dict[str, str] | None = None,
    chart_placeholders: list[str] | None = None,
    chart_images: list[Dict[str, Any]] | None = None,
    caution_text: str = DEFAULT_CAUTION_TEXT,
    title: str = "상담 요약 보고서",
    client_id: str = "",
    session_label: str = "",
    created_at: datetime | None = None,
) -> Dict[str, Any]:
    created_at = created_at or datetime.now()
    normalized_meta = dict(metadata or {})

    if client_id and "내담자 ID" not in normalized_meta:
        normalized_meta["내담자 ID"] = client_id

    if session_label and "회기" not in normalized_meta:
        normalized_meta["회기"] = session_label

    if "작성일" not in normalized_meta:
        normalized_meta["작성일"] = created_at.strftime("%Y-%m-%d")

    normalized_sections = dict(sections or {})

    if not normalized_sections:
        normalized_sections = _parse_markdown_sections(report_text)

    ordered_sections = []

    for key, section_title in REPORT_SECTION_ORDER:
        ordered_sections.append(
            (
                section_title,
                str(normalized_sections.get(key, "")).strip(),
            )
        )

    valid_chart_images = []

    for chart in chart_images or []:
        chart_title = str(chart.get("title", "첨부 차트")).strip() or "첨부 차트"
        image_bytes = chart.get("image_bytes")

        if not image_bytes:
            continue

        valid_chart_images.append(
            {
                "title": chart_title,
                "image_bytes": image_bytes,
            }
        )

    return {
        "title": title,
        "metadata": normalized_meta,
        "sections": ordered_sections,
        "chart_placeholders": chart_placeholders or DEFAULT_CHART_PLACEHOLDERS,
        "chart_images": valid_chart_images,
        "caution_text": caution_text or DEFAULT_CAUTION_TEXT,
    }


def _split_gender_age_region(meta: Dict[str, Any]) -> tuple[str, str]:
    gender_age_region = str(meta.get("성별/연령대/지역", "-"))
    gender_age = str(meta.get("성별·연령대", gender_age_region))
    region = str(meta.get("지역", "-"))

    if region == "-" and " · " in gender_age_region:
        parts = gender_age_region.split(" · ", 1)
        gender_age = parts[0].strip()
        region = parts[1].strip()

    return gender_age or "-", region or "-"


def make_docx_report_bytes(
    report_text: str,
    client_id: str = "",
    session_label: str = "",
    created_at: datetime | None = None,
    metadata: Dict[str, Any] | None = None,
    sections: Dict[str, str] | None = None,
    chart_placeholders: list[str] | None = None,
    chart_images: list[Dict[str, Any]] | None = None,
    caution_text: str = DEFAULT_CAUTION_TEXT,
    title: str = "상담 요약 보고서",
) -> bytes | None:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt, RGBColor
    except Exception:
        return None

    payload = _normalize_report_payload(
        report_text=report_text,
        metadata=metadata,
        sections=sections,
        chart_placeholders=chart_placeholders,
        chart_images=chart_images,
        caution_text=caution_text,
        title=title,
        client_id=client_id,
        session_label=session_label,
        created_at=created_at,
    )

    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    styles = doc.styles

    if "Normal" in styles:
        normal_style = styles["Normal"]
        normal_style.font.name = "맑은 고딕"
        normal_style.font.size = Pt(10.5)

    title_paragraph = doc.add_heading(payload["title"], level=0)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for run in title_paragraph.runs:
        run.font.name = "맑은 고딕"
        run.font.color.rgb = RGBColor(15, 23, 42)

    meta = payload["metadata"]
    gender_age, region = _split_gender_age_region(meta)

    meta_rows = [
        [
            "내담자 ID",
            meta.get("내담자 ID", client_id or "-"),
            "회기",
            meta.get("회기", session_label or "-"),
            "성별·연령대",
            gender_age or "-",
        ],
        [
            "지역",
            region or "-",
            "상담 분류",
            meta.get("상담 분류", "-"),
            "작성일",
            meta.get("작성일", "-"),
        ],
    ]

    table = doc.add_table(rows=2, cols=6)
    table.style = "Table Grid"

    for row_idx, row_values in enumerate(meta_rows):
        for col_idx, value in enumerate(row_values):
            cell = table.cell(row_idx, col_idx)
            cell.text = str(value or "-")

            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "맑은 고딕"
                    run.font.size = Pt(9.5)

                    if col_idx in (0, 2, 4):
                        run.bold = True

    doc.add_paragraph("")

    for index, (section_title, body) in enumerate(payload["sections"], start=1):
        heading = doc.add_heading(f"{index}. {section_title}", level=1)

        for run in heading.runs:
            run.font.name = "맑은 고딕"
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(29, 78, 216)

        paragraph = doc.add_paragraph(str(body or "-"))

        for run in paragraph.runs:
            run.font.name = "맑은 고딕"
            run.font.size = Pt(10.5)

        doc.add_paragraph("")

    chart_heading = doc.add_heading("6. 첨부 차트", level=1)

    for run in chart_heading.runs:
        run.font.name = "맑은 고딕"
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(29, 78, 216)

    if payload["chart_images"]:
        for chart in payload["chart_images"]:
            chart_title = str(chart.get("title", "첨부 차트"))
            image_bytes = chart.get("image_bytes")

            if not image_bytes:
                continue

            sub_heading = doc.add_heading(chart_title, level=2)

            for run in sub_heading.runs:
                run.font.name = "맑은 고딕"
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(15, 23, 42)

            try:
                doc.add_picture(BytesIO(image_bytes), width=Inches(6.2))
            except Exception:
                doc.add_paragraph(f"{chart_title} 이미지를 삽입하지 못했습니다.")

            doc.add_paragraph("")
    else:
        for placeholder in payload["chart_placeholders"]:
            doc.add_paragraph(str(placeholder))

    caution_heading = doc.add_heading("주의 문구", level=1)

    for run in caution_heading.runs:
        run.font.name = "맑은 고딕"
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(29, 78, 216)

    caution_paragraph = doc.add_paragraph(payload["caution_text"])

    for run in caution_paragraph.runs:
        run.font.name = "맑은 고딕"
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(30, 64, 175)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def make_pdf_report_bytes(
    report_text: str,
    client_id: str = "",
    session_label: str = "",
    created_at: datetime | None = None,
    metadata: Dict[str, Any] | None = None,
    sections: Dict[str, str] | None = None,
    chart_placeholders: list[str] | None = None,
    chart_images: list[Dict[str, Any]] | None = None,
    caution_text: str = DEFAULT_CAUTION_TEXT,
    title: str = "상담 요약 보고서",
) -> bytes | None:
    try:
        os.environ.setdefault("XDG_CACHE_HOME", "/tmp")
        os.environ.setdefault("DYLD_LIBRARY_PATH", "/opt/homebrew/lib")
        os.environ.setdefault("DYLD_FALLBACK_LIBRARY_PATH", "/opt/homebrew/lib")
        from weasyprint import HTML
    except Exception:
        return None

    payload = _normalize_report_payload(
        report_text=report_text,
        metadata=metadata,
        sections=sections,
        chart_placeholders=chart_placeholders,
        chart_images=chart_images,
        caution_text=caution_text,
        title=title,
        client_id=client_id,
        session_label=session_label,
        created_at=created_at,
    )

    meta = payload["metadata"]
    gender_age, region = _split_gender_age_region(meta)

    meta_rows = [
        [
            ("내담자 ID", meta.get("내담자 ID", client_id or "-")),
            ("회기", meta.get("회기", session_label or "-")),
            ("성별·연령대", gender_age or "-"),
        ],
        [
            ("지역", region or "-"),
            ("상담 분류", meta.get("상담 분류", "-")),
            ("작성일", meta.get("작성일", "-")),
        ],
    ]

    meta_html = "\n".join(
        "<tr>"
        + "".join(
            f"<th>{html_escape(str(label))}</th><td>{html_escape(str(value or '-'))}</td>"
            for label, value in row
        )
        + "</tr>"
        for row in meta_rows
    )

    section_html = "\n".join(
        f"""
        <section class="report-section">
            <h2>{index}. {html_escape(str(section_title))}</h2>
            <p>{html_escape(str(body or "-"))}</p>
        </section>
        """
        for index, (section_title, body) in enumerate(payload["sections"], start=1)
    )

    if payload["chart_images"]:
        chart_html = "\n".join(
            f"""
            <div class="chart-image-card">
                <h3>{html_escape(str(chart.get("title", "첨부 차트")))}</h3>
                <img src="data:image/png;base64,{base64.b64encode(chart.get("image_bytes", b"")).decode("utf-8")}" />
            </div>
            """
            for chart in payload["chart_images"]
            if chart.get("image_bytes")
        )
    else:
        chart_html = "\n".join(
            f"<div class='chart-placeholder'>{html_escape(str(placeholder))}</div>"
            for placeholder in payload["chart_placeholders"]
        )

    html = f"""
<!doctype html>
<html lang="ko">
<head>
    <meta charset="utf-8">
    <style>
        @page {{
            size: A4;
            margin: 18mm;
        }}

        body {{
            font-family: "Apple SD Gothic Neo", "Malgun Gothic", "Noto Sans CJK KR", sans-serif;
            color: #0F172A;
            font-size: 11pt;
            line-height: 1.75;
            word-break: keep-all;
            overflow-wrap: anywhere;
            background: #FFFFFF;
        }}

        .report-sheet {{
            padding: 0;
        }}

        .report-date {{
            text-align: right;
            color: #64748B;
            font-size: 9pt;
            margin-bottom: 3mm;
        }}

        .title {{
            color: #0F172A;
            font-size: 22pt;
            font-weight: 800;
            text-align: center;
            margin: 0 0 8mm;
        }}

        .meta-table {{
            width: 100%;
            border-collapse: separate;
            border-spacing: 0;
            border: 1px solid #D8E1F0;
            border-radius: 8px;
            overflow: hidden;
            margin: 0 0 9mm;
        }}

        .meta-table th,
        .meta-table td {{
            border-right: 1px solid #D8E1F0;
            border-bottom: 1px solid #D8E1F0;
            padding: 3mm 4mm;
            font-size: 10pt;
            text-align: left;
        }}

        .meta-table th {{
            background: #F8FAFC;
            color: #334155;
            font-weight: 700;
            width: 14%;
        }}

        .meta-table td {{
            background: #FFFFFF;
            color: #111827;
            font-weight: 600;
            width: 19%;
        }}

        .meta-table tr:last-child th,
        .meta-table tr:last-child td {{
            border-bottom: none;
        }}

        .meta-table th:last-child,
        .meta-table td:last-child {{
            border-right: none;
        }}

        .report-section {{
            margin: 0 0 7mm;
            page-break-inside: avoid;
        }}

        .report-section + .report-section {{
            border-top: 1px dashed #CBD5E1;
            padding-top: 5mm;
        }}

        .report-section h2 {{
            color: #1D4ED8;
            font-size: 13pt;
            margin: 0 0 2.5mm;
        }}

        .report-section p {{
            margin: 0;
            font-size: 10.5pt;
            line-height: 1.7;
            white-space: pre-wrap;
        }}

        .chart-report-section {{
            page-break-inside: auto;
        }}

        .chart-grid {{
            display: block;
        }}

        .chart-placeholder {{
            border: 1px dashed #CBD5E1;
            border-radius: 8px;
            padding: 8mm 4mm;
            margin-bottom: 5mm;
            text-align: center;
            color: #64748B;
            font-size: 9.5pt;
        }}

        .chart-image-card {{
            border: 1px solid #D8E1F0;
            border-radius: 10px;
            padding: 5mm;
            margin-bottom: 6mm;
            page-break-inside: avoid;
        }}

        .chart-image-card h3 {{
            margin: 0 0 4mm;
            color: #0F172A;
            font-size: 11pt;
            font-weight: 700;
        }}

        .chart-image-card img {{
            width: 100%;
            height: auto;
            display: block;
        }}

        .caution {{
            margin-top: 8mm;
            background: #EFF6FF;
            border: 1px solid #BFDBFE;
            border-radius: 8px;
            padding: 4mm;
            color: #1E40AF;
            font-size: 9.5pt;
            line-height: 1.6;
        }}
    </style>
</head>
<body>
    <div class="report-sheet">
        <div class="report-date">작성일: {html_escape(str(meta.get("작성일", "-")))}</div>
        <div class="title">{html_escape(payload["title"])}</div>

        <table class="meta-table">
            {meta_html}
        </table>

        {section_html}

        <section class="report-section chart-report-section">
            <h2>6. 첨부 차트</h2>
            <div class="chart-grid">
                {chart_html}
            </div>
        </section>

        <div class="caution">
            {html_escape(payload["caution_text"])}
        </div>
    </div>
</body>
</html>
"""

    return HTML(string=html).write_pdf()