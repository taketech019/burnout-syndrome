"""src/rag/ingest.py — RAG 인덱스 빌드 (ChromaDB + KoSBERT).

검색 소스:
1. AI Hub 심리상담 라벨링 데이터 (`data/references/심리상담데이터/Training/02.라벨링데이터/*.zip.part0`)
   → 각 zip 안의 JSON. paragraph[*].paragraph_text 추출.
2. 임상 가이드라인 PDF / 윤리 규정 docx (`data/references/*.pdf`, `*.docx`)
3. HIRA CSV (`data/raw/*.csv`) — 통계 컨텍스트 일부 추가 (header + 일부 행)

설계 결정:
- 전체 465K 건은 CPU 환경에서 임베딩에 시간이 너무 오래 걸려 MVP 데모용으로
  `MAX_AIHUB_DOCS` 한도(기본 3000건) 적용. 환경변수 또는 인자로 조정 가능.
- 발화는 너무 짧으면(< 20자) 스킵. 너무 길면(> 600자) 청킹.

빌드 명령: `python -m src.rag.ingest`
"""
import json
import logging
import os
import sys
import zipfile
from pathlib import Path
from typing import Iterator

from config import AIHUB_DIR, CHROMA_DIR, DATA_DIR, EMBEDDING_MODEL, RAW_DIR, REFERENCES_DIR

log = logging.getLogger("rag.ingest")
logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")


MAX_AIHUB_DOCS = int(os.getenv("RAG_MAX_AIHUB_DOCS", "3000"))
MAX_PDF_DOCS = int(os.getenv("RAG_MAX_PDF_DOCS", "500"))
MAX_CSV_DOCS = int(os.getenv("RAG_MAX_CSV_DOCS", "300"))
MIN_TEXT_LEN = 20
MAX_TEXT_LEN = 600


def _chunk(text: str, max_len: int = MAX_TEXT_LEN) -> list[str]:
    """긴 텍스트를 max_len 단위로 잘라 리스트로. 단어 단위로 자르지 않음 (한국어)."""
    text = text.strip()
    if len(text) <= max_len:
        return [text]
    out = []
    for i in range(0, len(text), max_len):
        out.append(text[i : i + max_len])
    return out


def _iter_aihub_paragraphs(root: Path) -> Iterator[dict]:
    """`data/references/심리상담데이터/Training/02.라벨링데이터/*.zip.part0`을
    하나씩 열어 paragraph_text를 발화 단위로 yield."""
    if not root.exists():
        log.warning("AI Hub 폴더 부재: %s", root)
        return
    candidates = []
    for sub in ("Training", "Validation"):
        labeling = root / sub / "02.라벨링데이터"
        if labeling.exists():
            candidates.extend(sorted(labeling.glob("*.zip*")))
    log.info("AI Hub ZIP 후보 %d개", len(candidates))

    for zpath in candidates:
        try:
            with zipfile.ZipFile(zpath) as zf:
                for name in zf.namelist():
                    if not name.endswith(".json"):
                        continue
                    try:
                        raw = zf.read(name)
                        data = json.loads(raw.decode("utf-8"))
                    except (UnicodeDecodeError, json.JSONDecodeError):
                        continue
                    paragraphs = data.get("paragraph") or data.get("paragraphs") or []
                    for i, p in enumerate(paragraphs):
                        if not isinstance(p, dict):
                            continue
                        speaker = str(p.get("paragraph_speaker") or p.get("speaker") or "").strip()
                        text = str(p.get("paragraph_text") or p.get("text") or "").strip()
                        if not text or len(text) < MIN_TEXT_LEN:
                            continue
                        for chunk in _chunk(text):
                            yield {
                                "text": (f"{speaker}: {chunk}" if speaker else chunk).strip(),
                                "source": f"AIHub/{zpath.name}::{name}#p{i}",
                                "type": "aihub_session",
                            }
        except zipfile.BadZipFile:
            log.warning("Bad zip skipped: %s", zpath)
            continue


def _iter_reference_pdfs(refs_dir: Path) -> Iterator[dict]:
    if not refs_dir.exists():
        return
    try:
        from pypdf import PdfReader
    except ImportError:
        log.warning("pypdf 미설치 — PDF skip")
        return
    for pdf_file in refs_dir.rglob("*.pdf"):
        try:
            reader = PdfReader(pdf_file)
        except Exception as e:
            log.warning("PDF 로드 실패 %s: %s", pdf_file.name, e)
            continue
        for i, page in enumerate(reader.pages):
            try:
                text = (page.extract_text() or "").strip()
            except Exception:
                continue
            if len(text) < MIN_TEXT_LEN:
                continue
            for chunk in _chunk(text, max_len=800):
                yield {
                    "text": chunk,
                    "source": f"{pdf_file.name}#p{i+1}",
                    "type": "clinical_guideline",
                }


def _iter_reference_docx(refs_dir: Path) -> Iterator[dict]:
    if not refs_dir.exists():
        return
    try:
        from docx import Document
    except ImportError:
        return
    for docx_file in refs_dir.rglob("*.docx"):
        try:
            doc = Document(docx_file)
        except Exception as e:
            log.warning("DOCX 로드 실패 %s: %s", docx_file.name, e)
            continue
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if len(text) < MIN_TEXT_LEN:
                continue
            for chunk in _chunk(text):
                yield {
                    "text": chunk,
                    "source": f"{docx_file.name}#p{i+1}",
                    "type": "ethics_or_reference",
                }


def _iter_hira_csv(raw_dir: Path) -> Iterator[dict]:
    """HIRA CSV의 헤더 + 일부 행을 자연어 풀어서 yield (RAG가 통계 질문에도 답하도록)."""
    if not raw_dir.exists():
        return
    try:
        import pandas as pd
    except ImportError:
        return
    for csv_file in raw_dir.glob("*.csv"):
        for enc in ("cp949", "utf-8-sig", "utf-8"):
            try:
                df = pd.read_csv(csv_file, encoding=enc, low_memory=False)
                break
            except (UnicodeDecodeError, UnicodeError):
                continue
            except Exception as e:
                log.warning("CSV 로드 실패 %s (%s): %s", csv_file.name, enc, e)
                df = None
                break
        else:
            log.warning("CSV 인코딩 미해결: %s", csv_file.name)
            continue
        if df is None or df.empty:
            continue
        cols = list(df.columns)
        header_text = f"[표] {csv_file.name} 열: {', '.join(map(str, cols))}"
        yield {"text": header_text, "source": f"HIRA/{csv_file.name}#header", "type": "hira_csv"}
        for i, row in df.head(MAX_CSV_DOCS).iterrows():
            sentence = ", ".join(f"{c}={row[c]}" for c in cols if str(row[c]).strip())
            if len(sentence) < MIN_TEXT_LEN:
                continue
            yield {
                "text": f"[{csv_file.stem}] {sentence}",
                "source": f"HIRA/{csv_file.name}#r{i}",
                "type": "hira_csv",
            }


def _make_embedder():
    """KoSBERT (sentence-transformers). 실패 시 LangChain SBERT."""
    from langchain_community.embeddings import HuggingFaceEmbeddings
    return HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL)


def build_index(
    max_aihub: int = MAX_AIHUB_DOCS,
    max_pdf: int = MAX_PDF_DOCS,
    max_csv: int = MAX_CSV_DOCS,
) -> dict:
    from langchain.schema import Document
    from langchain_community.vectorstores import Chroma

    log.info("KoSBERT 임베딩 로드 중...")
    embedder = _make_embedder()

    docs: list[Document] = []
    counts = {"aihub_session": 0, "clinical_guideline": 0, "ethics_or_reference": 0, "hira_csv": 0}

    log.info("AI Hub 라벨링 데이터 수집 (max=%d)", max_aihub)
    for item in _iter_aihub_paragraphs(AIHUB_DIR):
        docs.append(Document(page_content=item["text"], metadata={"source": item["source"], "type": item["type"]}))
        counts["aihub_session"] += 1
        if counts["aihub_session"] >= max_aihub:
            break
    log.info("AI Hub: %d개", counts["aihub_session"])

    log.info("PDF 가이드라인 수집 (max=%d)", max_pdf)
    for item in _iter_reference_pdfs(REFERENCES_DIR):
        docs.append(Document(page_content=item["text"], metadata={"source": item["source"], "type": item["type"]}))
        counts["clinical_guideline"] += 1
        if counts["clinical_guideline"] >= max_pdf:
            break
    log.info("PDF: %d개", counts["clinical_guideline"])

    log.info("DOCX 윤리규정 수집")
    for item in _iter_reference_docx(REFERENCES_DIR):
        docs.append(Document(page_content=item["text"], metadata={"source": item["source"], "type": item["type"]}))
        counts["ethics_or_reference"] += 1
    log.info("DOCX: %d개", counts["ethics_or_reference"])

    log.info("HIRA CSV 컨텍스트 수집 (max=%d/file)", max_csv)
    for item in _iter_hira_csv(RAW_DIR):
        docs.append(Document(page_content=item["text"], metadata={"source": item["source"], "type": item["type"]}))
        counts["hira_csv"] += 1
    log.info("HIRA: %d개", counts["hira_csv"])

    if not docs:
        return {"error": "인덱싱 대상 문서 0개. data/ 폴더 확인 필요.", **counts}

    CHROMA_DIR.mkdir(parents=True, exist_ok=True)
    log.info("ChromaDB 인덱싱 시작: 총 %d개 문서", len(docs))

    # 배치 임베딩으로 큰 컬렉션 처리.
    BATCH = 200
    store = None
    for start in range(0, len(docs), BATCH):
        batch = docs[start : start + BATCH]
        if store is None:
            store = Chroma.from_documents(
                batch,
                embedder,
                persist_directory=str(CHROMA_DIR),
                collection_name="counshelper",
            )
        else:
            store.add_documents(batch)
        log.info("  진행: %d / %d", min(start + BATCH, len(docs)), len(docs))

    # chromadb 0.4 호환 (persist 메서드). 0.5+ 에서는 no-op.
    try:
        store.persist()  # type: ignore[attr-defined]
    except (AttributeError, NotImplementedError):
        pass

    return {"total": len(docs), "persist_dir": str(CHROMA_DIR), **counts}


if __name__ == "__main__":
    print("RAG 인덱스 빌드 시작...")
    try:
        result = build_index()
    except Exception as e:
        print(f"빌드 중 오류: {e}")
        sys.exit(1)
    print(json.dumps(result, ensure_ascii=False, indent=2))
