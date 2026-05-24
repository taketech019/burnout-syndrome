"""scripts/scrape_mentalhealth.py

국가정신건강정보포털 50개 질환 문서 수집 → data/references/*.docx

확인된 HTML 구조 (1형 양극성장애 시범 수집으로 검증):
  div#tab1~#tab4 (탭 콘텐츠, DOM에 항상 존재, CSS on/off로만 숨김)
    div.box_sty06         → 인트로 본문
    ul.accordi.disease_info
      li (× N)
        div.tit           → 섹션 제목 (span.sr-only 제외)
        div.accordi_con   → 섹션 본문 (<br> 기반)

전략: JS로 모든 accordi_con을 display:block → #tab1~#tab4 직접 파싱
      탭 클릭 불필요 (콘텐츠가 DOM에 항상 있음)

실행:
    python scripts/scrape_mentalhealth.py
"""

import re
import urllib.parse
from io import BytesIO
from pathlib import Path
from urllib.parse import urljoin

import requests as req
from bs4 import BeautifulSoup, Tag
from docx import Document
from docx.shared import Inches
from playwright.sync_api import sync_playwright

DISEASES = [
    (25, "1형 양극성장애"),
    (24, "2형 양극성장애"),
    (18, "강박장애"),
    (73, "거식증"),
    (9,  "경계성 성격장애"),
    (23, "경도인지장애"),
    (33, "공황장애"),
    (55, "기면증"),
    (54, "노인우울증"),
    (61, "도박장애"),
    (57, "망상장애"),
    (63, "물질관련장애(마약성진통제)"),
    (72, "반사회성 인격장애"),
    (71, "배설장애"),
    (34, "범불안장애"),
    (31, "불면장애"),
    (35, "불안장애"),
    (68, "사회 불안장애"),
    (60, "섬망"),
    (10, "성격장애"),
    (15, "성인ADHD"),
    (32, "수면과 수면장애"),
    (70, "수면무호흡증"),
    (30, "스트레스"),
    (17, "신경성 폭식증"),
    (8,  "신체증상장애"),
    (22, "알츠하이머 치매"),
    (20, "알코올사용장애"),
    (64, "야경증"),
    (66, "여성우울증"),
    (65, "연극성 인격장애"),
    (28, "외상후 스트레스장애"),
    (38, "우울과 우울장애"),
    (67, "월경 전 불쾌감 증상"),
    (56, "일주기리듬 수면각성장애"),
    (19, "자살과 자살예방"),
    (14, "자폐스펙트럼장애"),
    (29, "적응장애"),
    (59, "전두측두엽 신경인지장애"),
    (27, "정신증"),
    (26, "조현병"),
    (36, "주요우울장애"),
    (16, "주의력결핍 과잉행동장애"),
    (21, "중독(의존)"),
    (62, "지적장애"),
    (12, "특정 학습장애"),
    (13, "틱장애"),
    (11, "품행장애"),
    (69, "하지 불안 증후군"),
    (58, "혈관성 신경인지장애"),
]

BASE_URL = "https://www.mentalhealth.go.kr/portal/disease/diseaseDetail.do"
LIST_URL = "https://www.mentalhealth.go.kr/portal/disease/diseaseList.do"

TAB_MAP = {
    "tab1": "개요",
    "tab2": "진단",
    "tab3": "치료",
    "tab4": "스스로 돕는 법",
}
TAB_NAMES = list(TAB_MAP.values())

OUT_DIR = Path("data/references")
OUT_DIR.mkdir(parents=True, exist_ok=True)

NOISE_RE = re.compile(
    r"^(접기|펼치기|열기|닫기|모두\s*(펼치기|접기)|사례.{0,6}(접기|펼치기)"
    r"|질환\s*정보\s*모두\s*(펼치기|접기)|이전|다음|목록|목록으로|top|▲|▼|위로"
    r"|자세히\s*보기|더\s*보기|닫\s*기|[0-9]+$)$",
    re.IGNORECASE,
)

IMG_SKIP_RE = re.compile(r"btn|icon|logo|arrow|bullet|blank|spacer|pixel|bg_", re.IGNORECASE)


# ── 파싱 유틸 ──────────────────────────────────────────────

def clean_text(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def extract_lines(el: Tag) -> list[str]:
    """
    엘리먼트에서 의미 있는 텍스트 라인 추출.
    이 사이트: <br> 기반 단락 + <table> 혼합.
    <br> → \\n 치환 후 줄 분리, 중복 제거.
    """
    el_copy = BeautifulSoup(str(el), "html.parser").find()
    if el_copy is None:
        return []
    for br in el_copy.find_all("br"):
        br.replace_with("\n")
    raw = el_copy.get_text("\n")
    seen: set[str] = set()
    lines: list[str] = []
    for seg in raw.split("\n"):
        t = clean_text(seg)
        if t and len(t) > 3 and not NOISE_RE.match(t) and t not in seen:
            lines.append(t)
            seen.add(t)
    return lines


def get_section_title(tit_div: Tag) -> str:
    """div.tit 에서 sr-only 스팬을 제외하고 제목 추출."""
    for sr in tit_div.find_all("span", class_="sr-only"):
        sr.decompose()
    return clean_text(tit_div.get_text())


def parse_tab(tab_div: Tag) -> dict[str, list[str]]:
    """
    탭 div → {섹션명: [본문줄...]}

    구조:
      div.box_sty06  → '__intro__' 섹션
      ul.accordi > li > div.tit + div.accordi_con
    """
    sections: dict[str, list[str]] = {}

    intro_div = tab_div.find("div", class_="box_sty06")
    if intro_div:
        lines = extract_lines(intro_div)
        if lines:
            sections["__intro__"] = lines

    for ul in tab_div.find_all("ul", class_="accordi"):
        for li in ul.find_all("li", recursive=False):
            tit_div = li.find("div", class_="tit")
            con_div = li.find("div", class_="accordi_con")
            if not tit_div:
                continue
            sec_name = get_section_title(tit_div)
            if not sec_name or NOISE_RE.match(sec_name):
                continue
            sections[sec_name] = extract_lines(con_div) if con_div else []

    return sections


def collect_images(tab_div: Tag, base_url: str) -> list[str]:
    urls: list[str] = []
    for img in tab_div.find_all("img", src=True):
        src = img["src"]
        if IMG_SKIP_RE.search(src):
            continue
        full = urljoin(base_url, src)
        if full not in urls:
            urls.append(full)
    return urls


def fetch_image(url: str, cookies: dict) -> bytes | None:
    try:
        r = req.get(url, cookies=cookies, timeout=10)
        if r.status_code == 200 and len(r.content) > 500:
            return r.content
    except Exception:
        pass
    return None


# ── docx 빌더 ──────────────────────────────────────────────

def build_docx(
    idx: int,
    diss_id: int,
    name: str,
    date: str,
    all_tabs: dict[str, dict],
    all_imgs: dict[str, list[str]],
    detail_url: str,
    cookies: dict,
) -> Path:
    doc = Document()
    doc.add_heading(name, level=0)

    meta = doc.add_paragraph()
    meta.add_run(
        f"dissId: {diss_id} | 감수일: {date or '미확인'}\n"
        f"출처: {detail_url}\n"
        f"발행기관: 보건복지부 국립정신건강센터, 대한신경정신의학회\n"
        f"라이선스: 공공누리 제4유형(출처표시·상업적이용금지·변경금지)\n"
        f"법적 한계 고지: 본 정보는 정신건강정보 이해를 돕는 참고자료이며, "
        f"개별 진단·치료는 의사 진료가 반드시 필요합니다."
    ).italic = True

    for tab_name in TAB_NAMES:
        doc.add_heading(tab_name, level=1)
        sections = all_tabs.get(tab_name, {})
        if not sections:
            doc.add_paragraph("[해당 탭 내용 없음]")
        else:
            for sec_name, lines in sections.items():
                if sec_name != "__intro__":
                    doc.add_heading(sec_name, level=2)
                for line in lines:
                    doc.add_paragraph(line)

        for img_url in all_imgs.get(tab_name, []):
            data = fetch_image(img_url, cookies)
            if data:
                try:
                    doc.add_picture(BytesIO(data), width=Inches(5.5))
                except Exception:
                    pass

    safe = name.replace(" ", "_").replace("/", "_").replace("(", "").replace(")", "")
    out = OUT_DIR / f"{idx:02d}_{safe}.docx"
    doc.save(out)
    return out


# ── 메인 ──────────────────────────────────────────────────

def scrape_one(page, diss_id: int, name: str, detail_url: str) -> tuple[dict, dict, str]:
    """단일 질환 수집. (all_tabs, all_imgs, date) 반환."""
    page.goto(detail_url, timeout=30_000)
    page.wait_for_load_state("networkidle", timeout=15_000)

    # 모든 아코디언 펼치기
    page.evaluate("""
        () => {
            document.querySelectorAll('.accordi_con').forEach(el => {
                el.style.display = 'block';
            });
            document.querySelectorAll('.tab_content').forEach(el => {
                el.style.display = 'block';
                el.classList.remove('off');
            });
        }
    """)
    page.wait_for_timeout(200)

    html = page.content()
    soup = BeautifulSoup(html, "html.parser")

    date = ""
    m = re.search(r"감수일\s*([0-9.]+)", soup.get_text())
    if m:
        date = m.group(1)

    all_tabs: dict[str, dict] = {}
    all_imgs: dict[str, list[str]] = {}

    for tab_id, tab_name in TAB_MAP.items():
        tab_div = soup.find("div", id=tab_id)
        if not tab_div:
            continue
        all_tabs[tab_name] = parse_tab(tab_div)
        all_imgs[tab_name] = collect_images(tab_div, detail_url)

    return all_tabs, all_imgs, date


def main():
    failed: list[tuple] = []

    with sync_playwright() as pw:
        browser = pw.chromium.launch(headless=True)
        context = browser.new_context(
            user_agent="Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
        )
        page = context.new_page()

        # 세션 확립
        page.goto(LIST_URL, timeout=30_000)
        page.wait_for_timeout(800)

        for idx, (diss_id, name) in enumerate(DISEASES, start=1):
            encoded = urllib.parse.quote(name)
            detail_url = f"{BASE_URL}?dissId={diss_id}&srCodeNm={encoded}"

            try:
                all_tabs, all_imgs, date = scrape_one(page, diss_id, name, detail_url)

                # 최소 콘텐츠 검증
                total_lines = sum(
                    sum(len(v) for v in secs.values())
                    for secs in all_tabs.values()
                )
                if total_lines < 5:
                    raise RuntimeError(f"콘텐츠 부족 ({total_lines}줄)")

                cookies = {c["name"]: c["value"] for c in context.cookies()}
                out = build_docx(idx, diss_id, name, date, all_tabs, all_imgs, detail_url, cookies)

                size_kb = out.stat().st_size // 1024
                tab_summary = " | ".join(
                    f"{t}:{sum(len(v) for v in all_tabs.get(t, {}).values())}줄"
                    for t in TAB_NAMES
                )
                print(f"[{idx:02d}/50] OK  {name}  ({size_kb}KB) [{tab_summary}]")

            except Exception as e:
                print(f"[{idx:02d}/50] FAIL {name}: {e}")
                failed.append((diss_id, name, str(e)))

            page.wait_for_timeout(1200)

        browser.close()

    if failed:
        fail_path = Path("failed.txt")
        with fail_path.open("w", encoding="utf-8") as f:
            for did, nm, err in failed:
                f.write(f"{did}\t{nm}\t{err}\n")
        print(f"\n실패 {len(failed)}건 -> failed.txt 확인")

    total = len(list(OUT_DIR.glob("[0-9][0-9]_*.docx")))
    print(f"\ndata/references/ 저장 파일: {total}개")


if __name__ == "__main__":
    main()
